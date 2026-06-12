from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


out = Path("20260318-mr-migration-completion-cn.docx")

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

styles = doc.styles
styles["Normal"].font.name = "SimSun"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.line_spacing = 1.25
styles["Normal"].paragraph_format.space_after = Pt(6)

for style_name in ["Title", "Heading 1", "Heading 2"]:
    style = styles[style_name]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("媒体关系部")
run.bold = True
p.add_run("\n电话：+41-44-234 85 00")

p = doc.add_paragraph("瑞银瑞士股份公司，新闻稿，2026年3月18日，第1页，共1页")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p = doc.add_paragraph("2026年3月18日")
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

h = doc.add_paragraph()
r = h.add_run("新闻稿")
r.bold = True
r.font.size = Pt(14)

h = doc.add_paragraph()
r = h.add_run("瑞银成功完成瑞士客户迁移，成为整合进程中的又一关键里程碑")
r.bold = True
r.font.size = Pt(16)

paragraphs = [
    "苏黎世，2026年3月18日 - 瑞银今日宣布，随着在瑞士入账客户迁移工作的完成，瑞银已在全球范围内成功完成所有原瑞士信贷客户向瑞银基础设施的迁移，这是瑞士信贷整合过程中的又一个关键里程碑。",
    "集团首席执行官Sergio P. Ermotti表示：“随着瑞士客户迁移工作的成功完成，我们已在全球范围内顺利迁移约120万名客户。凭借同事们的坚定决心、勤勉努力和投入，我们在两家全球系统重要性银行首次合并、以及银行业历史上最复杂的整合之一中，又达成了一个关键里程碑。要完成整合仍有大量工作要做，但客户迁移的结束强化了瑞银的业务实力，并为向所有客户提供更广泛、更顺畅的服务奠定了基础。”",
    "自2023年3月收购瑞士信贷以来，整合工作一直按照精心排序并执行的计划推进，并在2024年夏季达成了多个重要里程碑，包括母行合并以及瑞士实体合并。瑞士迁移工作的准备包括扩大分行、联络中心和支持职能的能力，开展超过80,000次测试，并为一线同事提供超过132,000小时的迁移专项培训。此外，瑞银平台上的支付量增加了25%，达到每日近310万笔交易，体现出新整合系统的稳健性和效率。",
    "此次迁移高度重视客户旅程，并在各客户群体中收到积极反馈。瑞银在整个过程中保持透明沟通，发送了近300万封个性化信函，并建立专门的数字中心，用于客户信息发布和自助服务。因此，客户满意度保持在高水平，信任度有所提升，反映出瑞银持续坚持客户至上的承诺。",
    "客户迁移工作的完成使整合进入最后阶段，包括停用旧有IT基础设施。正如此前宣布的那样，瑞银正按计划在2026年底前基本完成整合。",
]

for text in paragraphs:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p = doc.add_paragraph()
p.add_run("瑞银瑞士股份公司").bold = True

doc.add_paragraph("瑞士媒体关系部：+41 44 234 85 00")
doc.add_paragraph("www.ubs.com/media")

p = doc.add_paragraph()
r = p.add_run("注：本文件为中文翻译稿，依据原英文新闻稿翻译。")
r.italic = True
r.font.size = Pt(9)

for para in doc.paragraphs:
    for run in para.runs:
        run.font.name = "SimSun"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

doc.save(out)
print(out.resolve())
